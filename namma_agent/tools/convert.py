"""Document conversion — turn Markdown into the file format the user actually asked for.

The agent writes Markdown natively (LLMs are good at it). When a user explicitly
wants a *different* format — "give me a Word doc", "export as PDF", "make it a
PowerPoint" — this tool converts the Markdown into that format and drops a
downloadable file into the chat.

  convert_document(content|path, to) — md → docx/pdf/pptx/html/txt/odt/rtf/epub/rst/tex/…

Engine strategy (graceful degradation, the project pattern):
  * If ``pandoc`` is on PATH it handles every format at high fidelity.
  * Without pandoc, self-contained fallbacks still cover the common asks:
      - ``txt``  : strip Markdown to plain text          (stdlib)
      - ``html`` : Markdown → standalone HTML            (stdlib mini-renderer)
      - ``docx`` : Markdown → Word                       (python-docx, already a dep)
  * Anything else without pandoc returns a clear "install pandoc" message rather
    than crashing.

Output files land in ``data/media/documents/`` and are surfaced inline as a
download link (same ``data.url`` convention the media tools use).
"""
from __future__ import annotations

import html as _html
import re
import shutil
import subprocess
import uuid
from pathlib import Path

from namma_agent.core.logger import logger
from namma_agent.core.safety import check_path
from namma_agent.core.tools import ToolRegistry, ToolResult

_OUT_DIR = Path("data/media/documents")

# Friendly names → canonical extension.
_ALIASES = {
    "word": "docx", "doc": "docx", "msword": "docx",
    "powerpoint": "pptx", "ppt": "pptx", "slides": "pptx", "deck": "pptx",
    "markdown": "md", "htm": "html", "webpage": "html",
    "text": "txt", "plain": "txt", "plaintext": "txt",
    "latex": "tex", "openoffice": "odt", "libreoffice": "odt",
    "ebook": "epub", "restructuredtext": "rst",
}

# Formats the self-contained fallbacks can produce with no external binary.
_BUILTIN = {"md", "txt", "html", "docx"}

# ext → pandoc writer name (when it differs from the extension).
_PANDOC_WRITER = {"tex": "latex", "txt": "plain", "adoc": "asciidoc"}
# Writers that benefit from a self-contained standalone document.
_STANDALONE = {"html", "epub", "tex", "rtf"}


def _normalize_format(to: str) -> str:
    to = (to or "").strip().lower().lstrip(".")
    return _ALIASES.get(to, to)


def _slugify(text: str) -> str:
    s = re.sub(r"[^a-z0-9]+", "-", (text or "").strip().lower()).strip("-")
    return (s or "document")[:60]


def _resolve_markdown(args: dict) -> tuple[str, str]:
    """Return (markdown_text, base_name). Source is inline ``content`` or a ``path``."""
    content = args.get("content")
    if content and content.strip():
        return content, ""
    path = (args.get("path") or "").strip()
    if not path:
        raise ValueError("provide either 'content' (Markdown text) or 'path' (a .md file)")
    ok, reason = check_path(path)
    if not ok:
        raise ValueError(reason)
    p = Path(path).expanduser()
    if not p.is_file():
        raise ValueError(f"not a file: {path}")
    return p.read_text(encoding="utf-8", errors="replace"), p.stem


# ── Block parser (shared by the HTML + DOCX fallbacks) ───────────────────────
#
# A deliberately small Markdown subset — headings, paragraphs, bullet/ordered
# lists, fenced code, block-quotes, pipe tables, and horizontal rules. Good
# enough for the documents the agent produces; pandoc covers the long tail.

def _parse_blocks(md: str) -> list[tuple]:
    blocks: list[tuple] = []
    lines = md.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    i, n = 0, len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Fenced code block.
        fence = re.match(r"^\s*```(.*)$", line)
        if fence:
            lang = fence.group(1).strip()
            body: list[str] = []
            i += 1
            while i < n and not re.match(r"^\s*```", lines[i]):
                body.append(lines[i])
                i += 1
            i += 1  # skip closing fence
            blocks.append(("code", lang, "\n".join(body)))
            continue

        # Heading.
        h = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if h:
            blocks.append(("heading", len(h.group(1)), h.group(2).strip()))
            i += 1
            continue

        # Horizontal rule.
        if re.match(r"^([-*_])\s*(\1\s*){2,}$", stripped):
            blocks.append(("hr",))
            i += 1
            continue

        # Pipe table: a header row followed by a |---|---| separator.
        if "|" in stripped and i + 1 < n and re.match(r"^\s*\|?[\s:|-]+\|?\s*$", lines[i + 1]) \
                and "-" in lines[i + 1]:
            headers = _split_row(stripped)
            rows = []
            i += 2
            while i < n and "|" in lines[i].strip() and lines[i].strip():
                rows.append(_split_row(lines[i].strip()))
                i += 1
            blocks.append(("table", headers, rows))
            continue

        # Block-quote (collapse consecutive > lines).
        if stripped.startswith(">"):
            quote = []
            while i < n and lines[i].strip().startswith(">"):
                quote.append(re.sub(r"^\s*>\s?", "", lines[i]))
                i += 1
            blocks.append(("quote", " ".join(q.strip() for q in quote)))
            continue

        # Bullet / ordered list (collapse consecutive item lines).
        if re.match(r"^\s*([-*+]|\d+[.)])\s+", line):
            ordered = bool(re.match(r"^\s*\d+[.)]\s+", line))
            items = []
            while i < n and re.match(r"^\s*([-*+]|\d+[.)])\s+", lines[i]):
                items.append(re.sub(r"^\s*([-*+]|\d+[.)])\s+", "", lines[i]).strip())
                i += 1
            blocks.append(("list", ordered, items))
            continue

        # Paragraph — gather until a blank line or a block starter.
        para = [stripped]
        i += 1
        while i < n and lines[i].strip() and not re.match(
                r"^\s*(#{1,6}\s|```|>|[-*+]\s|\d+[.)]\s)", lines[i]):
            para.append(lines[i].strip())
            i += 1
        blocks.append(("para", " ".join(para)))
    return blocks


def _split_row(row: str) -> list[str]:
    return [c.strip() for c in row.strip().strip("|").split("|")]


_INLINE_RE = re.compile(
    r"(\*\*.+?\*\*|__.+?__|\*.+?\*|_.+?_|`.+?`|\[.+?\]\(.+?\))")


def _iter_inline(text: str):
    """Yield (kind, payload) spans: text / bold / italic / code / (link_text, href)."""
    for part in _INLINE_RE.split(text):
        if not part:
            continue
        if (part.startswith("**") and part.endswith("**")) or \
           (part.startswith("__") and part.endswith("__")):
            yield "bold", part[2:-2]
        elif part.startswith("`") and part.endswith("`"):
            yield "code", part[1:-1]
        elif (part.startswith("*") and part.endswith("*")) or \
             (part.startswith("_") and part.endswith("_")):
            yield "italic", part[1:-1]
        else:
            m = re.match(r"^\[(.+?)\]\((.+?)\)$", part)
            if m:
                yield "link", (m.group(1), m.group(2))
            else:
                yield "text", part


# ── txt fallback ─────────────────────────────────────────────────────────────

def _md_to_txt(md: str) -> str:
    out: list[str] = []
    for block in _parse_blocks(md):
        kind = block[0]
        if kind == "heading":
            out.append(_strip_inline(block[2]).upper())
        elif kind == "para":
            out.append(_strip_inline(block[1]))
        elif kind == "list":
            for j, item in enumerate(block[2], 1):
                bullet = f"{j}." if block[1] else "-"
                out.append(f"  {bullet} {_strip_inline(item)}")
        elif kind == "code":
            out.append(block[2])
        elif kind == "quote":
            out.append(f"> {_strip_inline(block[1])}")
        elif kind == "table":
            out.append("  ".join(block[1]))
            for r in block[2]:
                out.append("  ".join(r))
        elif kind == "hr":
            out.append("-" * 40)
        out.append("")
    return "\n".join(out).strip() + "\n"


def _strip_inline(text: str) -> str:
    return "".join(
        payload if kind != "link" else payload[0]
        for kind, payload in _iter_inline(text))


# ── html fallback ────────────────────────────────────────────────────────────

_HTML_SHELL = """<!DOCTYPE html>
<html lang="en"><head><meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>{title}</title>
<style>
 body{{font:16px/1.6 -apple-system,Segoe UI,Roboto,sans-serif;max-width:46rem;
   margin:2.5rem auto;padding:0 1.25rem;color:#1a1a1a}}
 h1,h2,h3,h4{{line-height:1.25;margin:1.6em 0 .5em}}
 code{{background:#f3f3f3;padding:.1em .35em;border-radius:4px;font-size:.92em}}
 pre{{background:#f6f8fa;padding:1rem;border-radius:8px;overflow:auto}}
 pre code{{background:none;padding:0}}
 blockquote{{border-left:4px solid #ddd;margin:1em 0;padding:.2em 1em;color:#555}}
 table{{border-collapse:collapse;width:100%;margin:1em 0}}
 th,td{{border:1px solid #ddd;padding:.5em .7em;text-align:left}}
 th{{background:#f6f8fa}}
 hr{{border:none;border-top:1px solid #ddd;margin:2em 0}}
</style></head><body>
{body}
</body></html>
"""


def _inline_to_html(text: str) -> str:
    out = []
    for kind, payload in _iter_inline(text):
        if kind == "bold":
            out.append(f"<strong>{_html.escape(payload)}</strong>")
        elif kind == "italic":
            out.append(f"<em>{_html.escape(payload)}</em>")
        elif kind == "code":
            out.append(f"<code>{_html.escape(payload)}</code>")
        elif kind == "link":
            out.append(f'<a href="{_html.escape(payload[1])}">{_html.escape(payload[0])}</a>')
        else:
            out.append(_html.escape(payload))
    return "".join(out)


def _md_to_html(md: str, title: str) -> str:
    # Prefer the `markdown` library when present — richer + battle-tested.
    try:
        import markdown as _md_lib  # noqa: PLC0415

        body = _md_lib.markdown(md, extensions=["tables", "fenced_code"])
        return _HTML_SHELL.format(title=_html.escape(title), body=body)
    except ImportError:
        pass

    parts: list[str] = []
    for block in _parse_blocks(md):
        kind = block[0]
        if kind == "heading":
            lvl = min(block[1], 6)
            parts.append(f"<h{lvl}>{_inline_to_html(block[2])}</h{lvl}>")
        elif kind == "para":
            parts.append(f"<p>{_inline_to_html(block[1])}</p>")
        elif kind == "list":
            tag = "ol" if block[1] else "ul"
            items = "".join(f"<li>{_inline_to_html(it)}</li>" for it in block[2])
            parts.append(f"<{tag}>{items}</{tag}>")
        elif kind == "code":
            parts.append(f"<pre><code>{_html.escape(block[2])}</code></pre>")
        elif kind == "quote":
            parts.append(f"<blockquote>{_inline_to_html(block[1])}</blockquote>")
        elif kind == "table":
            head = "".join(f"<th>{_inline_to_html(h)}</th>" for h in block[1])
            rows = "".join(
                "<tr>" + "".join(f"<td>{_inline_to_html(c)}</td>" for c in r) + "</tr>"
                for r in block[2])
            parts.append(f"<table><thead><tr>{head}</tr></thead><tbody>{rows}</tbody></table>")
        elif kind == "hr":
            parts.append("<hr>")
    return _HTML_SHELL.format(title=_html.escape(title), body="\n".join(parts))


# ── docx fallback (python-docx) ──────────────────────────────────────────────

def _add_inline_runs(paragraph, text: str) -> None:
    for kind, payload in _iter_inline(text):
        if kind == "link":
            run = paragraph.add_run(payload[0])
            run.underline = True
        else:
            run = paragraph.add_run(payload)
            if kind == "bold":
                run.bold = True
            elif kind == "italic":
                run.italic = True
            elif kind == "code":
                run.font.name = "Consolas"


def _md_to_docx(md: str, out_path: Path, title: str) -> None:
    try:
        import docx  # noqa: PLC0415
        from docx.shared import Pt  # noqa: PLC0415
    except ImportError as exc:
        raise RuntimeError(
            "writing .docx needs the 'python-docx' package (pip install python-docx) "
            "— or install 'pandoc' for richer conversion") from exc

    doc = docx.Document()
    for block in _parse_blocks(md):
        kind = block[0]
        if kind == "heading":
            doc.add_heading(block[2], level=min(block[1], 6))
        elif kind == "para":
            _add_inline_runs(doc.add_paragraph(), block[1])
        elif kind == "list":
            style = "List Number" if block[1] else "List Bullet"
            for item in block[2]:
                _add_inline_runs(doc.add_paragraph(style=style), item)
        elif kind == "code":
            p = doc.add_paragraph()
            run = p.add_run(block[2])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        elif kind == "quote":
            p = doc.add_paragraph(style="Intense Quote")
            _add_inline_runs(p, block[1])
        elif kind == "table" and block[1]:
            table = doc.add_table(rows=1, cols=len(block[1]))
            table.style = "Table Grid"
            for cell, text in zip(table.rows[0].cells, block[1]):
                cell.paragraphs[0].add_run(text).bold = True
            for row in block[2]:
                cells = table.add_row().cells
                for cell, text in zip(cells, row):
                    _add_inline_runs(cell.paragraphs[0], text)
        elif kind == "hr":
            doc.add_paragraph("_" * 40)
    doc.save(str(out_path))


# ── pandoc (universal, high fidelity) ────────────────────────────────────────

def _pandoc(md: str, out_path: Path, fmt: str, title: str) -> None:
    pandoc = shutil.which("pandoc")
    if not pandoc:
        raise RuntimeError("pandoc not found")
    cmd = [pandoc, "-f", "gfm", "-o", str(out_path)]
    writer = _PANDOC_WRITER.get(fmt)
    if writer:
        cmd += ["-t", writer]
    if fmt in _STANDALONE:
        cmd.append("--standalone")
    if title:
        cmd += ["--metadata", f"title={title}"]
    proc = subprocess.run(cmd, input=md, capture_output=True, text=True,
                          encoding="utf-8", errors="replace", timeout=120)
    if proc.returncode != 0 or not out_path.exists():
        raise RuntimeError((proc.stderr or "pandoc failed").strip()[:400])


# ── Handler ──────────────────────────────────────────────────────────────────

def _convert(args: dict) -> ToolResult:
    fmt = _normalize_format(args.get("to") or args.get("format") or "")
    if not fmt:
        return ToolResult(ok=False, content="",
                          error="'to' (target format, e.g. docx/pdf/pptx/html/txt) is required")
    try:
        md, base_name = _resolve_markdown(args)
    except ValueError as exc:
        return ToolResult(ok=False, content="", error=str(exc))

    # Title: explicit → first heading → source filename.
    title = (args.get("title") or "").strip()
    if not title:
        m = re.search(r"^#{1,6}\s+(.+)$", md, re.MULTILINE)
        title = (m.group(1).strip() if m else base_name) or "Document"

    base = _slugify(args.get("filename") or title or base_name)
    _OUT_DIR.mkdir(parents=True, exist_ok=True)
    out_path = _OUT_DIR / f"{base}-{uuid.uuid4().hex[:8]}.{fmt}"

    have_pandoc = shutil.which("pandoc") is not None
    try:
        if fmt == "md":
            out_path.write_text(md, encoding="utf-8")
        elif have_pandoc and fmt not in {"md"}:
            # Pandoc wins on fidelity whenever it's available.
            _pandoc(md, out_path, fmt, title)
        elif fmt == "txt":
            out_path.write_text(_md_to_txt(md), encoding="utf-8")
        elif fmt == "html":
            out_path.write_text(_md_to_html(md, title), encoding="utf-8")
        elif fmt == "docx":
            _md_to_docx(md, out_path, title)
        else:
            return ToolResult(
                ok=False, content="",
                error=f"converting to .{fmt} needs 'pandoc' on PATH (https://pandoc.org/install). "
                      f"Without it I can produce: md, txt, html, docx.")
    except Exception as exc:  # noqa: BLE001
        logger.warning("[convert] %s → %s failed: %s", base_name or "markdown", fmt, exc)
        return ToolResult(ok=False, content="", error=f"conversion to .{fmt} failed: {exc}")

    size = out_path.stat().st_size
    url = f"/api/media/documents/{out_path.name}"
    filename = out_path.name
    content = f"📄 **{title}** → `.{fmt}` ({_human_size(size)}) — [⬇ Download {filename}]({url})"
    logger.info("[convert] wrote %s (%d bytes)", out_path, size)
    return ToolResult(ok=True, content=content,
                      data={"url": url, "kind": "document", "format": fmt, "bytes": size})


def _human_size(n: int) -> str:
    if n < 1024:
        return f"{n} B"
    if n < 1024 * 1024:
        return f"{n / 1024:.1f} KB"
    return f"{n / (1024 * 1024):.1f} MB"


def register(registry: ToolRegistry) -> None:
    registry.register(
        "convert_document",
        "Convert Markdown into a downloadable file in the format the user asked for "
        "(docx/Word, pdf, pptx/PowerPoint, html, txt, odt, rtf, epub, rst, tex, …). "
        "Use this when the user explicitly wants a non-Markdown file. Pass the Markdown "
        "as 'content' (or a .md file 'path') and the target as 'to'.",
        {
            "type": "object",
            "properties": {
                "content": {"type": "string",
                            "description": "the Markdown text to convert (preferred)"},
                "path": {"type": "string",
                         "description": "path to an existing .md file (alternative to 'content')"},
                "to": {"type": "string",
                       "description": "target format: docx, pdf, pptx, html, txt, odt, rtf, "
                                      "epub, rst, tex, md"},
                "title": {"type": "string",
                          "description": "document title (defaults to the first heading)"},
                "filename": {"type": "string",
                             "description": "base name for the output file (optional)"},
            },
            "required": ["to"],
        },
        _convert)
